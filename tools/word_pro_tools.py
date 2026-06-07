"""
word_pro_tools.py — Professional Word (.docx) authoring (python-docx).

Beyond create/read/find-replace: headings, formatted paragraphs, tables with
styling, images, bullet/numbered lists, page breaks, headers/footers, and a
RTL/Arabic-aware paragraph helper. No Microsoft Office required.

Tools:
  • word_add_heading(path, text, level)            — heading (level 0=title … 4)
  • word_add_paragraph(path, text, ...)            — styled paragraph (bold/size/align/color)
  • word_add_table(path, data, style, header_bold) — table from JSON, styled
  • word_add_image(path, image_path, width_inches) — insert an image
  • word_add_list(path, items, numbered)           — bullet or numbered list
  • word_add_page_break(path)                      — start a new page
  • word_set_header_footer(path, header, footer)   — document header/footer text
  • word_set_rtl(path)                             — make the document right-to-left (Arabic)
"""
from __future__ import annotations

import json
from pathlib import Path

from langchain_core.tools import tool


# ── helpers ───────────────────────────────────────────────────────────────────
def _require_docx():
    try:
        import docx  # noqa: F401
        return None
    except ImportError:
        return "❌ python-docx غير مثبّت. ثبّته بـ: pip install python-docx"


def _open_or_create(path: str):
    """Open an existing .docx, or create a new one if missing."""
    from docx import Document
    p = Path(path)
    if p.is_file():
        try:
            return Document(str(p)), p, None
        except Exception as e:
            return None, None, f"❌ تعذّر فتح المستند: {e}"
    p.parent.mkdir(parents=True, exist_ok=True)
    return Document(), p, None


def _rgb(hex_str: str):
    from docx.shared import RGBColor
    h = (hex_str or "").lstrip("#").strip()
    if len(h) != 6:
        return None
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ═══════════════════════════════════════════════════════════════════════════════
@tool
def word_add_heading(path: str, text: str, level: int = 1) -> str:
    """إضافة عنوان إلى مستند Word (يُنشئ المستند إن لم يكن موجوداً).

    Args:
        path: مسار المستند
        text: نص العنوان
        level: 0 = عنوان رئيسي للمستند، 1..4 = عناوين فرعية بأحجام متدرجة
    """
    err = _require_docx()
    if err:
        return err
    doc, p, err = _open_or_create(path)
    if err:
        return err
    try:
        doc.add_heading(text, level=max(0, min(level, 4)))
        doc.save(str(p))
        return f"✅ أُضيف عنوان (مستوى {level}): «{text}»"
    except Exception as e:
        return f"❌ خطأ: {e}"


@tool
def word_add_paragraph(
    path: str,
    text: str,
    bold: bool = False,
    italic: bool = False,
    size: int = 0,
    align: str = "",
    color: str = "",
) -> str:
    """إضافة فقرة نصية بتنسيق (غامق/مائل/حجم/محاذاة/لون).

    Args:
        path: مسار المستند
        text: نص الفقرة
        bold/italic: غامق/مائل
        size: حجم الخط بالنقاط (0 = افتراضي)
        align: 'center'/'right'/'left'/'justify'
        color: لون الخط hex مثل '1F4E79'
    """
    err = _require_docx()
    if err:
        return err
    doc, p, err = _open_or_create(path)
    if err:
        return err
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color and _rgb(color):
            run.font.color.rgb = _rgb(color)
        align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if align in align_map:
            para.alignment = align_map[align]
        doc.save(str(p))
        return f"✅ أُضيفت فقرة ({len(text)} حرف)."
    except Exception as e:
        return f"❌ خطأ: {e}"


@tool
def word_add_table(
    path: str,
    data: str,
    style: str = "Light Grid Accent 1",
    header_bold: bool = True,
) -> str:
    """إضافة جدول إلى مستند Word من بيانات JSON.

    Args:
        path: مسار المستند
        data: JSON — قائمة قوائم (أول صف عناوين) أو قائمة قواميس.
              مثال: [["البند","المبلغ"],["إيجار",3000],["رواتب",12000]]
        style: نمط الجدول، مثل 'Light Grid Accent 1' أو 'Medium Shading 1 Accent 1'
               أو 'Table Grid' (حدود بسيطة)
        header_bold: تغميق صف العناوين
    """
    err = _require_docx()
    if err:
        return err
    doc, p, err = _open_or_create(path)
    if err:
        return err
    try:
        rows_data = json.loads(data) if isinstance(data, str) else data
        if not rows_data:
            return "❌ لا توجد بيانات للجدول."
        if isinstance(rows_data[0], dict):
            headers = list(rows_data[0].keys())
            table_rows = [headers] + [[r.get(h, "") for h in headers] for r in rows_data]
        else:
            table_rows = [list(r) for r in rows_data]

        n_cols = len(table_rows[0])
        table = doc.add_table(rows=0, cols=n_cols)
        try:
            table.style = style
        except Exception:
            table.style = "Table Grid"  # safe fallback

        for r, row in enumerate(table_rows):
            cells = table.add_row().cells
            for c in range(n_cols):
                val = row[c] if c < len(row) else ""
                cells[c].text = str(val)
                if r == 0 and header_bold:
                    for para in cells[c].paragraphs:
                        for run in para.runs:
                            run.bold = True
        doc.save(str(p))
        return f"✅ أُضيف جدول ({len(table_rows)}×{n_cols}) بنمط «{table.style.name}»."
    except json.JSONDecodeError:
        return "❌ بيانات JSON غير صحيحة."
    except Exception as e:
        return f"❌ خطأ في الجدول: {e}"


@tool
def word_add_image(path: str, image_path: str, width_inches: float = 5.0) -> str:
    """إدراج صورة في المستند بعرض محدد (بالبوصة)."""
    err = _require_docx()
    if err:
        return err
    if not Path(image_path).is_file():
        return f"❌ الصورة غير موجودة: {image_path}"
    doc, p, err = _open_or_create(path)
    if err:
        return err
    try:
        from docx.shared import Inches
        doc.add_picture(image_path, width=Inches(width_inches))
        doc.save(str(p))
        return f"✅ أُدرجت الصورة بعرض {width_inches} بوصة."
    except Exception as e:
        return f"❌ خطأ في إدراج الصورة: {e}"


@tool
def word_add_list(path: str, items: str, numbered: bool = False) -> str:
    """إضافة قائمة نقطية أو رقمية.

    Args:
        path: مسار المستند
        items: عناصر القائمة، عنصر لكل سطر
        numbered: True للقائمة الرقمية، False للنقطية
    """
    err = _require_docx()
    if err:
        return err
    doc, p, err = _open_or_create(path)
    if err:
        return err
    try:
        style = "List Number" if numbered else "List Bullet"
        lines = [l.strip() for l in items.split("\n") if l.strip()]
        for line in lines:
            try:
                doc.add_paragraph(line, style=style)
            except Exception:
                # fallback if the named style is missing in the template
                doc.add_paragraph(("• " if not numbered else "") + line)
        doc.save(str(p))
        return f"✅ أُضيفت قائمة {'رقمية' if numbered else 'نقطية'} ({len(lines)} عنصر)."
    except Exception as e:
        return f"❌ خطأ: {e}"


@tool
def word_add_page_break(path: str) -> str:
    """إضافة فاصل صفحة (بدء صفحة جديدة)."""
    err = _require_docx()
    if err:
        return err
    doc, p, err = _open_or_create(path)
    if err:
        return err
    try:
        doc.add_page_break()
        doc.save(str(p))
        return "✅ أُضيف فاصل صفحة."
    except Exception as e:
        return f"❌ خطأ: {e}"


@tool
def word_set_header_footer(path: str, header: str = "", footer: str = "") -> str:
    """تعيين نص رأس وتذييل الصفحة لكل المستند."""
    err = _require_docx()
    if err:
        return err
    doc, p, err = _open_or_create(path)
    if err:
        return err
    try:
        section = doc.sections[0]
        if header:
            section.header.paragraphs[0].text = header
        if footer:
            section.footer.paragraphs[0].text = footer
        doc.save(str(p))
        parts = []
        if header:
            parts.append("رأس")
        if footer:
            parts.append("تذييل")
        return f"✅ تم تعيين {' و'.join(parts) or 'الرأس/التذييل'}."
    except Exception as e:
        return f"❌ خطأ: {e}"


@tool
def word_set_rtl(path: str) -> str:
    """جعل المستند من اليمين إلى اليسار (مناسب للعربية) لكل الفقرات."""
    err = _require_docx()
    if err:
        return err
    doc, p, err = _open_or_create(path)
    if err:
        return err
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        count = 0
        for para in doc.paragraphs:
            pPr = para._p.get_or_add_pPr()
            bidi = pPr.find(qn("w:bidi"))
            if bidi is None:
                bidi = OxmlElement("w:bidi")
                pPr.append(bidi)
            count += 1
        doc.save(str(p))
        return f"✅ تم ضبط اتجاه {count} فقرة من اليمين لليسار (RTL)."
    except Exception as e:
        return f"❌ خطأ: {e}"
