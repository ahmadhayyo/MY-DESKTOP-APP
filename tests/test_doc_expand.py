"""
Document-expansion tests — run with:  python -m unittest tests.test_doc_expand

The model is mocked so tests are deterministic and free. They lock in the
guarantees that the earlier failed attempt violated: tables are read in body
order, never dropped, and a failed/greedy model response falls back to the
original blocks so content is never lost.
"""
import json
import os
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core import ai_office as ai       # noqa: E402
from core import doc_expand as de      # noqa: E402


def _make_doc(path):
    """A small doc with interleaved headings, paragraphs and a table."""
    from docx import Document
    d = Document()
    d.add_heading("المحور الثامن", level=1)
    d.add_paragraph("مقدمة قصيرة عن التصميم.")
    t = d.add_table(rows=2, cols=3)
    hdr = ["القسم", "الموضوع", "المدة"]
    for j, v in enumerate(hdr):
        t.rows[0].cells[j].text = v
    for j, v in enumerate(["1", "الألوان", "120"]):
        t.rows[1].cells[j].text = v
    d.add_heading("الجلسة 1", level=2)
    d.add_paragraph("هدف الجلسة.")
    d.save(path)


class Reader(unittest.TestCase):
    def setUp(self):
        self.d = tempfile.mkdtemp()
        self.src = os.path.join(self.d, "in.docx")
        _make_doc(self.src)

    def test_reads_blocks_in_order_with_table(self):
        blocks = de.read_docx_blocks(self.src)
        types = [b["type"] for b in blocks]
        self.assertEqual(types, ["heading", "para", "table", "heading", "para"])
        tbl = blocks[2]
        self.assertEqual(tbl["header"], ["القسم", "الموضوع", "المدة"])
        self.assertEqual(tbl["rows"], [["1", "الألوان", "120"]])

    def test_heading_levels(self):
        blocks = de.read_docx_blocks(self.src)
        self.assertEqual(blocks[0]["level"], 1)
        self.assertEqual(blocks[3]["level"], 2)


class Segmenting(unittest.TestCase):
    def test_new_section_per_top_heading(self):
        blocks = [
            {"type": "heading", "level": 1, "text": "A"},
            {"type": "para", "text": "x"},
            {"type": "heading", "level": 1, "text": "B"},
            {"type": "para", "text": "y"},
        ]
        secs = de.segment_sections(blocks)
        self.assertEqual(len(secs), 2)

    def test_large_section_splits(self):
        big = [{"type": "heading", "level": 1, "text": "A"}]
        big += [{"type": "para", "text": "z" * 1000} for _ in range(10)]
        secs = de.segment_sections(big, max_chars=2000)
        self.assertGreater(len(secs), 1)


class ExpandGuarantees(unittest.TestCase):
    def tearDown(self):
        ai._LLM_INVOKER = None

    def test_expansion_enriches(self):
        section = [{"type": "para", "text": "قصير"}]
        ai._LLM_INVOKER = lambda m: json.dumps(
            {"blocks": [{"type": "para", "text": "نص أطول وأكثر تفصيلاً واحترافية."}]},
            ensure_ascii=False)
        out = de.expand_section(section)
        self.assertIn("احترافية", out[0]["text"])

    def test_dropped_table_falls_back_to_original(self):
        section = [
            {"type": "table", "header": ["ح"], "rows": [["1"]]},
            {"type": "para", "text": "n"},
        ]
        # model returns text but DROPS the table → must keep original section
        ai._LLM_INVOKER = lambda m: json.dumps(
            {"blocks": [{"type": "para", "text": "expanded"}]}, ensure_ascii=False)
        out = de.expand_section(section)
        self.assertEqual(out, section)   # unchanged: table preserved

    def test_bad_json_falls_back(self):
        section = [{"type": "para", "text": "keep"}]
        ai._LLM_INVOKER = lambda m: "not json at all"
        self.assertEqual(de.expand_section(section), section)

    def test_model_error_falls_back(self):
        section = [{"type": "para", "text": "keep"}]
        def boom(m):
            raise RuntimeError("model down")
        ai._LLM_INVOKER = boom
        self.assertEqual(de.expand_section(section), section)


class RenderAndOrchestrate(unittest.TestCase):
    def setUp(self):
        self.d = tempfile.mkdtemp()
        self.src = os.path.join(self.d, "in.docx")
        _make_doc(self.src)

    def tearDown(self):
        ai._LLM_INVOKER = None

    def test_render_roundtrip_keeps_table(self):
        blocks = de.read_docx_blocks(self.src)
        out = os.path.join(self.d, "out.docx")
        de.render_docx(blocks, out, title="نسخة")
        back = de.read_docx_blocks(out)
        self.assertEqual(sum(1 for b in back if b["type"] == "table"), 1)

    def test_expand_document_never_loses_tables(self):
        # a model that always drops tables → orchestrator must still keep all
        ai._LLM_INVOKER = lambda m: json.dumps(
            {"blocks": [{"type": "para", "text": "expanded only"}]}, ensure_ascii=False)
        out = os.path.join(self.d, "exp.docx")
        stats = de.expand_document(self.src, out)
        self.assertEqual(stats["tables_out"], stats["tables_in"])
        self.assertEqual(stats["tables_out"], 1)
        back = de.read_docx_blocks(out)
        self.assertEqual(sum(1 for b in back if b["type"] == "table"), 1)

    def test_expand_document_grows_words(self):
        # a model that expands paragraphs (and echoes tables) → words increase
        def expander(messages):
            import re
            # echo any table blocks; lengthen paragraphs
            human = messages[-1].content
            data = json.loads(human.split("SECTION:\n", 1)[1].split("\n\nإضافة", 1)[0])
            out = []
            for b in data["blocks"]:
                if b["type"] == "para":
                    out.append({"type": "para", "text": b["text"] + " تفاصيل إضافية مهنية دقيقة."})
                else:
                    out.append(b)
            return json.dumps({"blocks": out}, ensure_ascii=False)
        ai._LLM_INVOKER = expander
        out = os.path.join(self.d, "grow.docx")
        stats = de.expand_document(self.src, out)
        self.assertGreater(stats["words_out"], stats["words_in"])
        self.assertEqual(stats["tables_out"], stats["tables_in"])


if __name__ == "__main__":
    unittest.main(verbosity=2)
