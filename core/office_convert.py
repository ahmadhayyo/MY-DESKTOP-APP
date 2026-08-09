"""
core/office_convert.py — the precise Office data & conversion engine.

Built for a specific person: someone who does data entry in Microsoft Office,
in Arabic *and* English, and loses nights retyping data into Excel by hand. This
module turns that night of work into one call, and converts precisely between
Excel · Word · PDF · PowerPoint · CSV.

Design
------
Everything here is pure/deterministic where it can be, so it is unit-tested:
  • Text intelligence — delimiter detection, Arabic-digit normalisation, header
    detection, numeric coercion, RTL detection, Arabic shaping for PDF.
  • Format I/O — read/write rows for xlsx / csv / docx-tables / pdf-tables, and
    render rows into a *formatted* Excel (bold frozen header, autofit, borders,
    RTL sheet when Arabic), a Word table, or a PDF table with correct Arabic.
  • `convert(src, target, dest)` — one dispatcher that picks the right precise
    engine per format pair, with a LibreOffice fallback for anything exotic.

Only pure-Python libraries are used (openpyxl, python-docx, pypdf, reportlab,
python-pptx, pdfplumber, arabic_reshaper, python-bidi) so it works on a plain
Windows machine with no LibreOffice required for the common Office pairs.
"""
from __future__ import annotations

import csv
import io
import os
import re
from pathlib import Path

# ── Arabic / digits helpers ───────────────────────────────────────────────────
# Arabic-Indic (٠-٩) and Extended/Persian (۰-۹) → ASCII digits.
_DIGIT_MAP = str.maketrans("٠١٢٣٤٥٦٧٨٩۰۱۲۳۴۵۶۷۸۹", "01234567890123456789")
_AR_RE = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]")
# Arabic decimal separator (٫) and thousands (٬)
_AR_DECIMAL = "\u066b"
_AR_THOUSANDS = "\u066c"


def normalize_digits(s):
    """Map Arabic-Indic/Persian digits (and separators) to ASCII. Non-str passes
    through untouched."""
    if not isinstance(s, str):
        return s
    return (s.translate(_DIGIT_MAP)
             .replace(_AR_THOUSANDS, ",")
             .replace(_AR_DECIMAL, "."))


def is_arabic(s) -> bool:
    return bool(isinstance(s, str) and _AR_RE.search(s))


def any_arabic(rows) -> bool:
    for row in rows:
        for cell in row:
            if is_arabic(cell):
                return True
    return False


def shape_arabic(s: str) -> str:
    """Reshape + bidi-reorder Arabic so it renders correctly in reportlab PDF
    (which has no native Arabic shaping). English text is returned unchanged."""
    if not is_arabic(s):
        return s
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        return get_display(arabic_reshaper.reshape(s))
    except Exception:
        return s


# ── numeric coercion ──────────────────────────────────────────────────────────
_NUM_RE = re.compile(r"^[+-]?\d{1,3}(,\d{3})*(\.\d+)?$|^[+-]?\d+(\.\d+)?$")


def coerce_value(v):
    """Turn a numeric-looking string into int/float (after digit normalisation),
    so Excel treats it as a number, not text. Everything else is stripped text."""
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return v
    s = normalize_digits(str(v)).strip()
    if s == "":
        return ""
    core = s.replace(",", "")
    if _NUM_RE.match(s):
        try:
            if "." in core:
                return float(core)
            return int(core)
        except ValueError:
            return s
    # percentage like "12%" → keep as text (Excel would need a format); leave clean
    return s


# ── delimiter detection & tabular parsing ─────────────────────────────────────
_CANDIDATE_DELIMS = ["\t", ",", ";", "\u060c", "|"]  # tab, comma, semicolon, ، , pipe


def detect_delimiter(text: str) -> str:
    """Pick the delimiter that splits the most lines into a *consistent* column
    count. Falls back to runs of 2+ spaces (common in pasted/aligned data)."""
    lines = [ln for ln in text.replace("\r\n", "\n").split("\n") if ln.strip()]
    sample = lines[:50]
    best, best_score = None, 0.0
    for d in _CANDIDATE_DELIMS:
        counts = [ln.count(d) for ln in sample]
        present = [c for c in counts if c > 0]
        if not present:
            continue
        # reward: many lines have it AND the column count is consistent
        from collections import Counter
        mode_c, mode_n = Counter(present).most_common(1)[0]
        score = mode_n * (mode_c + 1)  # consistency × width
        if score > best_score:
            best, best_score = d, score
    if best is not None:
        return best
    # multi-space fallback only if lines actually have aligned gaps
    if any(re.search(r"\S {2,}\S", ln) for ln in sample):
        return "  "  # sentinel meaning "2+ spaces regex split"
    return ""  # single column


def _split_line(line: str, delim: str) -> list[str]:
    if delim == "  ":
        return [c.strip() for c in re.split(r" {2,}|\t", line.strip())]
    if delim == "":
        return [line.strip()]
    if delim == ",":
        # honour quotes for real CSV
        return [c.strip() for c in next(csv.reader([line]))] if line else [""]
    return [c.strip() for c in line.split(delim)]


def _looks_like_header(first: list[str], rest: list[list[str]]) -> bool:
    """Heuristic: header row is non-numeric while the body has numbers below it."""
    if not rest:
        return False
    def numericish(cells):
        vals = [c for c in cells if c != ""]
        if not vals:
            return 0.0
        n = sum(1 for c in vals if _NUM_RE.match(normalize_digits(c).strip()))
        return n / len(vals)
    head_num = numericish(first)
    body_num = sum(numericish(r) for r in rest[:10]) / min(len(rest), 10)
    # Header row reads as labels (little/no numbers) while the body below carries
    # numbers. One numeric column is enough of a signal.
    return head_num < 0.2 and body_num > head_num + 0.05


def parse_tabular(text: str, has_header: str = "auto") -> dict:
    """Parse pasted/raw text into a clean rectangular table.

    Returns {"rows": list[list[str]], "header": bool, "delimiter": str}.
    Rows are rag-right padded to equal width; cells are trimmed and Arabic
    digits normalised in-place for display consistency.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    raw_lines = [ln for ln in text.split("\n") if ln.strip() != ""]
    if not raw_lines:
        return {"rows": [], "header": False, "delimiter": ""}
    delim = detect_delimiter(text)
    rows = [_split_line(ln, delim) for ln in raw_lines]
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    if has_header == "auto":
        header = _looks_like_header(rows[0], rows[1:])
    else:
        header = str(has_header).lower() in ("1", "true", "yes", "y")
    return {"rows": rows, "header": header, "delimiter": delim}


# ── Excel writing (formatted) ─────────────────────────────────────────────────
def rows_to_xlsx(rows, path: str, sheet_name: str = "Sheet1", header: bool = True,
                 rtl=None, title: str = "", coerce: bool = True) -> str:
    """Write rows into a nicely-formatted .xlsx: bold frozen header, thin borders,
    auto-fit column widths, numeric cells coerced, RTL sheet when Arabic."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment

    if rtl is None:
        rtl = any_arabic(rows)

    wb = Workbook()
    ws = wb.active
    ws.title = (sheet_name or "Sheet1")[:31]
    ws.sheet_view.rightToLeft = bool(rtl)

    start = 1
    if title:
        ws.cell(row=1, column=1, value=title).font = Font(bold=True, size=14)
        start = 3

    thin = Side(style="thin", color="D0D0D0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    head_fill = PatternFill("solid", fgColor="2E75B6")
    head_font = Font(bold=True, color="FFFFFF")
    align = Alignment(horizontal="right" if rtl else "left", vertical="center",
                      wrap_text=False)

    widths: dict[int, int] = {}
    for i, row in enumerate(rows):
        r = start + i
        is_head = header and i == 0
        for j, val in enumerate(row, start=1):
            cell = ws.cell(row=r, column=j,
                           value=(val if (is_head or not coerce) else coerce_value(val)))
            cell.border = border
            cell.alignment = align
            if is_head:
                cell.fill = head_fill
                cell.font = head_font
            widths[j] = max(widths.get(j, 8), min(60, len(str(val)) + 2))
    for j, w in widths.items():
        ws.column_dimensions[ws.cell(row=start, column=j).column_letter].width = w
    if header and rows:
        ws.freeze_panes = ws.cell(row=start + 1, column=1)

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out))
    return str(out)


# ── readers → rows ────────────────────────────────────────────────────────────
def xlsx_to_rows(path: str, sheet: str | None = None) -> list[list]:
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True)
    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append(["" if c is None else c for c in row])
    # drop wholly-empty trailing rows
    while rows and all(c == "" for c in rows[-1]):
        rows.pop()
    return rows


def csv_to_rows(path: str) -> list[list[str]]:
    with open(path, "r", encoding="utf-8-sig", newline="") as fh:
        sample = fh.read(4096)
        fh.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        return [list(r) for r in csv.reader(fh, dialect)]


def rows_to_csv(rows, path: str) -> str:
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    # utf-8-sig so Excel opens Arabic CSV without mojibake
    with open(out, "w", encoding="utf-8-sig", newline="") as fh:
        w = csv.writer(fh)
        for row in rows:
            w.writerow(list(row))
    return str(out)


def docx_tables_to_rows(path: str) -> list[list[list[str]]]:
    """Return every table in a Word doc as a list of row-matrices."""
    from docx import Document
    doc = Document(path)
    tables = []
    for tbl in doc.tables:
        matrix = []
        for row in tbl.rows:
            matrix.append([cell.text.strip() for cell in row.cells])
        if matrix:
            tables.append(matrix)
    return tables


def pdf_tables_to_rows(path: str) -> list[list[list[str]]]:
    """Extract tables from a PDF using pdfplumber (falls back to text lines)."""
    tables = []
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                for tbl in page.extract_tables() or []:
                    matrix = [["" if c is None else str(c).strip() for c in row]
                              for row in tbl]
                    if matrix:
                        tables.append(matrix)
    except Exception:
        pass
    return tables


# ── writers from rows ─────────────────────────────────────────────────────────
def rows_to_docx_table(rows, path: str, header: bool = True, rtl=None,
                       title: str = "") -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if rtl is None:
        rtl = any_arabic(rows)
    doc = Document()
    if title:
        h = doc.add_heading(title, level=1)
        if rtl:
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if not rows:
        doc.save(path)
        return str(path)
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Light Grid Accent 1"
    if rtl:
        # right-to-left table direction
        from docx.oxml.ns import qn
        tblPr = table._tbl.tblPr
        bidi = tblPr.makeelement(qn("w:bidiVisual"), {})
        tblPr.append(bidi)
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(cols):
            txt = str(row[j]) if j < len(row) else ""
            cells[j].text = txt
            for p in cells[j].paragraphs:
                if rtl:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(11)
                    if header and i == 0:
                        run.font.bold = True
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


_AR_FONT_REGISTERED = False


def _register_arabic_font():
    """Register a Windows TTF that supports Arabic for reportlab. Returns the
    font name to use ('HAYOAr' if Arabic-capable, else 'Helvetica')."""
    global _AR_FONT_REGISTERED
    if _AR_FONT_REGISTERED:
        return "HAYOAr"
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    for cand in (r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\tahoma.ttf",
                 r"C:\Windows\Fonts\segoeui.ttf",
                 "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"):
        if os.path.exists(cand):
            try:
                pdfmetrics.registerFont(TTFont("HAYOAr", cand))
                _AR_FONT_REGISTERED = True
                return "HAYOAr"
            except Exception:
                continue
    return "Helvetica"


def rows_to_pdf_table(rows, path: str, header: bool = True, rtl=None,
                      title: str = "") -> str:
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    if rtl is None:
        rtl = any_arabic(rows)
    font = _register_arabic_font()

    def cell(v):
        return shape_arabic(str(v)) if rtl else str(v)

    data = [[cell(c) for c in row] for row in rows]
    if rtl:
        data = [list(reversed(r)) for r in data]  # visual column order for RTL

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(out), pagesize=landscape(A4),
                            leftMargin=24, rightMargin=24, topMargin=28, bottomMargin=24)
    story = []
    if title:
        styles = getSampleStyleSheet()
        ts = ParagraphStyle("t", parent=styles["Title"], fontName=font,
                            alignment=(2 if rtl else 0))
        story.append(Paragraph(shape_arabic(title) if rtl else title, ts))
        story.append(Spacer(1, 10))
    if data:
        tbl = Table(data, repeatRows=1 if header else 0)
        style = [
            ("FONTNAME", (0, 0), (-1, -1), font),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0B0B0")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (-1, -1), "RIGHT" if rtl else "LEFT"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F6FB")]),
        ]
        if header:
            style += [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E75B6")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
            ]
        tbl.setStyle(TableStyle(style))
        story.append(tbl)
    doc.build(story)
    return str(out)


# ── the dispatcher ────────────────────────────────────────────────────────────
_TABULAR_IN = {".xlsx", ".xlsm", ".csv", ".tsv", ".txt", ".docx", ".pdf"}


def read_any_rows(src: str) -> list[list]:
    """Best-effort: read the first/most meaningful table of rows from any source."""
    ext = Path(src).suffix.lower()
    if ext in (".xlsx", ".xlsm"):
        return xlsx_to_rows(src)
    if ext in (".csv", ".tsv"):
        return csv_to_rows(src)
    if ext == ".txt":
        return parse_tabular(Path(src).read_text(encoding="utf-8", errors="replace"))["rows"]
    if ext == ".docx":
        tabs = docx_tables_to_rows(src)
        return tabs[0] if tabs else []
    if ext == ".pdf":
        tabs = pdf_tables_to_rows(src)
        return tabs[0] if tabs else []
    raise ValueError(f"صيغة غير مدعومة للقراءة: {ext}")


def _win32_to_pdf(src: str, dest: str) -> bool:
    """Highest-fidelity Word/PowerPoint → PDF using the user's installed
    Microsoft Office via COM (Windows only). Returns True on success. This gives
    perfect Arabic/RTL rendering because it's real Office doing the export."""
    if os.name != "nt":
        return False
    ext = Path(src).suffix.lower()
    src_abs, dest_abs = str(Path(src).resolve()), str(Path(dest).resolve())
    try:
        import win32com.client  # type: ignore
    except Exception:
        return False
    try:
        if ext in (".docx", ".doc", ".rtf"):
            word = win32com.client.Dispatch("Word.Application")
            try:
                doc = word.Documents.Open(src_abs, ReadOnly=True)
                doc.SaveAs(dest_abs, FileFormat=17)  # 17 = wdFormatPDF
                doc.Close(False)
                return os.path.exists(dest_abs)
            finally:
                word.Quit()
        if ext in (".pptx", ".ppt"):
            ppt = win32com.client.Dispatch("PowerPoint.Application")
            try:
                deck = ppt.Presentations.Open(src_abs, WithWindow=False)
                deck.SaveAs(dest_abs, 32)  # 32 = ppSaveAsPDF
                deck.Close()
                return os.path.exists(dest_abs)
            finally:
                ppt.Quit()
        if ext in (".xlsx", ".xlsm", ".xls"):
            excel = win32com.client.Dispatch("Excel.Application")
            try:
                wb = excel.Workbooks.Open(src_abs, ReadOnly=True)
                wb.ExportAsFixedFormat(0, dest_abs)  # 0 = xlTypePDF
                wb.Close(False)
                return os.path.exists(dest_abs)
            finally:
                excel.Quit()
    except Exception:
        return False
    return False


def _libreoffice_convert(src: str, target: str, dest: str) -> bool:
    """Fidelity fallback via LibreOffice/soffice if it's installed. Returns
    True on success. Used only for pairs the pure-Python paths don't cover."""
    import shutil, subprocess
    exe = shutil.which("soffice") or shutil.which("libreoffice")
    if not exe:
        return False
    outdir = str(Path(dest).parent)
    try:
        subprocess.run([exe, "--headless", "--convert-to", target, "--outdir",
                        outdir, src], capture_output=True, timeout=180)
        produced = Path(outdir) / (Path(src).stem + f".{target}")
        if produced.exists():
            if str(produced) != str(dest):
                produced.replace(dest)
            return True
    except Exception:
        pass
    return False


def convert(src: str, target_format: str, dest: str | None = None) -> str:
    """Convert `src` to `target_format` precisely. Returns the output path.

    Covered pure-Python pairs (both Arabic & English):
      xlsx→csv/pdf/docx · csv/txt→xlsx · docx→xlsx/pdf · pdf→xlsx · pptx→pdf
    Anything else attempts a LibreOffice fallback.
    """
    src = os.path.expanduser(src)
    if not os.path.exists(src):
        raise FileNotFoundError(src)
    ext = Path(src).suffix.lower()
    tgt = target_format.lstrip(".").lower()
    if dest is None:
        dest = str(Path(src).with_suffix(f".{tgt}"))
    dest = os.path.expanduser(dest)
    Path(dest).parent.mkdir(parents=True, exist_ok=True)

    # ── tabular pairs ─────────────────────────────────────────────
    if ext in (".xlsx", ".xlsm") and tgt == "csv":
        return rows_to_csv(xlsx_to_rows(src), dest)
    if ext in (".csv", ".tsv", ".txt") and tgt in ("xlsx", "xls"):
        if ext in (".csv", ".tsv"):
            rows, header = csv_to_rows(src), True
        else:
            parsed = parse_tabular(Path(src).read_text(encoding="utf-8", errors="replace"))
            rows, header = parsed["rows"], parsed["header"]
        return rows_to_xlsx(rows, dest, header=header)
    if ext in (".xlsx", ".xlsm") and tgt == "pdf":
        # Prefer real Excel (keeps her formatting); fall back to a clean table.
        if _win32_to_pdf(src, dest):
            return dest
        return rows_to_pdf_table(xlsx_to_rows(src), dest)
    if ext in (".xlsx", ".xlsm") and tgt in ("docx", "doc"):
        return rows_to_docx_table(xlsx_to_rows(src), dest)
    if ext == ".docx" and tgt in ("xlsx", "xls"):
        tabs = docx_tables_to_rows(src)
        if not tabs:
            raise ValueError("لا يوجد جدول في ملف Word لتحويله إلى Excel.")
        # merge all tables stacked (blank row between)
        merged = []
        for i, t in enumerate(tabs):
            if i:
                merged.append([""])
            merged += t
        return rows_to_xlsx(merged, dest, header=True)
    if ext == ".pdf" and tgt in ("xlsx", "xls"):
        tabs = pdf_tables_to_rows(src)
        if not tabs:
            raise ValueError("تعذّر العثور على جداول في ملف PDF.")
        merged = []
        for i, t in enumerate(tabs):
            if i:
                merged.append([""])
            merged += t
        return rows_to_xlsx(merged, dest, header=True)

    # ── document → pdf (fidelity: MS Office COM → LibreOffice → reportlab) ──
    if tgt == "pdf" and ext in (".docx", ".pptx", ".doc", ".ppt"):
        if _win32_to_pdf(src, dest):
            return dest
        if _libreoffice_convert(src, "pdf", dest):
            return dest
        if ext == ".docx":  # basic text fallback
            from docx import Document
            doc = Document(src)
            lines = [p.text for p in doc.paragraphs]
            rtl = any(is_arabic(x) for x in lines)
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            font = _register_arabic_font()
            d = SimpleDocTemplate(dest, pagesize=A4)
            st = getSampleStyleSheet()
            body = ParagraphStyle("b", parent=st["Normal"], fontName=font,
                                  alignment=(2 if rtl else 0), fontSize=11, leading=16)
            story = []
            for ln in lines:
                if ln.strip():
                    story.append(Paragraph(shape_arabic(ln) if rtl else ln, body))
                else:
                    story.append(Spacer(1, 6))
            d.build(story)
            return dest
        raise ValueError(f"تحويل {ext}→pdf يحتاج LibreOffice (غير مثبّت).")

    # ── last resort ──
    if _libreoffice_convert(src, tgt, dest):
        return dest
    raise ValueError(f"زوج التحويل غير مدعوم مباشرةً: {ext} → {tgt}")


if __name__ == "__main__":  # tiny smoke test
    import tempfile
    d = tempfile.mkdtemp()
    raw = "الاسم, العمر, المدينة\nأحمد, ٣٥, دمشق\nليلى, ٢٨, حلب"
    p = parse_tabular(raw)
    print("delim:", repr(p["delimiter"]), "header:", p["header"], "rows:", p["rows"])
    xlsx = rows_to_xlsx(p["rows"], os.path.join(d, "t.xlsx"), header=p["header"])
    print("xlsx:", xlsx, os.path.getsize(xlsx), "bytes")
    print("roundtrip rows:", xlsx_to_rows(xlsx))
    pdf = rows_to_pdf_table(p["rows"], os.path.join(d, "t.pdf"))
    print("pdf:", pdf, os.path.getsize(pdf), "bytes")
    print("office_convert smoke OK")
