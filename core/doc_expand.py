"""
core/doc_expand.py — expand a Word study/curriculum precisely and professionally.

Why this exists
---------------
A humanitarian training curriculum (World Vision) had to be expanded to be more
precise and professional. The earlier attempt failed badly: it flattened the
document to paragraph text, DROPPED all 30 tables (sessions, timings, learning
objectives), then stalled. This module fixes every one of those failures:

  1. It reads the .docx in true body order, keeping paragraphs AND tables — so
     nothing is ever lost.
  2. It segments the document into bounded sections (by heading), so each model
     call is small enough to be precise and reliable.
  3. It expands each section with a strong model, under a strict JSON contract
     (enrich, never delete; keep every table; keep Arabic), and — crucially — if
     any section's model call fails or returns malformed output, it FALLS BACK to
     that section's ORIGINAL blocks. The result can be richer, but never poorer.
  4. It renders a clean RTL Word document, headings/paragraphs/bullets/tables.

Block schema (the lingua franca between reader, model and renderer):
    {"type": "heading", "level": int, "text": str}
    {"type": "para",    "text": str}
    {"type": "bullet",  "text": str}
    {"type": "table",   "header": [str,...], "rows": [[cell,...],...]}
"""
from __future__ import annotations

import json

from core import ai_office as _ai


# ── read .docx in body order (paragraphs + tables) ────────────────────────────
def iter_block_items(doc):
    """Yield Paragraph and Table objects in the order they appear in the body.
    python-docx exposes paragraphs and tables separately and loses their
    interleaving; this walks the XML body to preserve true document order."""
    from docx.document import Document as _Doc
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent = doc.element.body
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _heading_level(par) -> int | None:
    """Return 0 for Title, 1..9 for Heading N, else None."""
    try:
        name = (par.style.name or "").strip()
    except Exception:
        return None
    if name.lower() == "title":
        return 0
    if name.lower().startswith("heading"):
        digits = "".join(ch for ch in name if ch.isdigit())
        return int(digits) if digits else 1
    return None


def read_docx_blocks(path: str) -> list[dict]:
    """Read a .docx into an ordered list of block dicts (paragraphs + tables)."""
    from docx import Document
    doc = Document(path)
    blocks: list[dict] = []
    for item in iter_block_items(doc):
        cls = item.__class__.__name__
        if cls == "Paragraph":
            text = (item.text or "").strip()
            if not text:
                continue
            lvl = _heading_level(item)
            if lvl is not None:
                blocks.append({"type": "heading", "level": lvl, "text": text})
            else:
                style = ""
                try:
                    style = (item.style.name or "").lower()
                except Exception:
                    style = ""
                if "list" in style or "bullet" in style:
                    blocks.append({"type": "bullet", "text": text})
                else:
                    blocks.append({"type": "para", "text": text})
        elif cls == "Table":
            rows = []
            for row in item.rows:
                rows.append([c.text.strip() for c in row.cells])
            if rows:
                header = rows[0]
                blocks.append({"type": "table", "header": header, "rows": rows[1:]})
    return blocks


# ── segment into bounded sections ─────────────────────────────────────────────
def segment_sections(blocks: list[dict], max_chars: int = 4500) -> list[list[dict]]:
    """Group blocks into sections. A new top-level heading (level<=2) starts a
    new section; a section that grows past max_chars is also split, so every
    model call stays small and precise. Tables count toward the size."""
    sections: list[list[dict]] = []
    cur: list[dict] = []
    cur_len = 0

    def block_len(b: dict) -> int:
        if b["type"] == "table":
            return sum(len(str(c)) for r in [b.get("header", [])] + b.get("rows", []) for c in r)
        return len(b.get("text", ""))

    for b in blocks:
        is_heading = b["type"] == "heading" and b.get("level", 9) <= 2
        if cur and (is_heading or cur_len >= max_chars):
            sections.append(cur)
            cur, cur_len = [], 0
        cur.append(b)
        cur_len += block_len(b)
    if cur:
        sections.append(cur)
    return sections


# ── expand one section with the strong model ──────────────────────────────────
_EXPAND_SYS = (
    "You are a senior instructional designer producing professional humanitarian "
    "vocational-training curricula (World Vision, North-Syria camps), in Arabic. "
    "You receive a SECTION of a training guide as a JSON list of blocks. Expand it "
    "to be more precise, detailed and professional WITHOUT losing anything.\n"
    "Return ONLY a JSON object: { \"blocks\": [ ...same block schema... ] }.\n"
    "Block schema: {\"type\":\"heading\",\"level\":int,\"text\":str} | "
    "{\"type\":\"para\",\"text\":str} | {\"type\":\"bullet\",\"text\":str} | "
    "{\"type\":\"table\",\"header\":[str],\"rows\":[[str]]}.\n"
    "Rules — critical:\n"
    "- NEVER drop a table. Keep every table and its data; you MAY add columns/rows "
    "or clarify cells, but never remove information.\n"
    "- Keep all headings; keep Arabic in Arabic; preserve numbers and timings.\n"
    "- Enrich paragraphs with concrete, professional detail: clearer learning "
    "objectives, step-by-step facilitation, materials, timing, assessment, and "
    "practical tips relevant to camp settings. Do NOT invent false facts.\n"
    "- Output valid JSON only, no prose, no code fences."
)


def expand_section(section: list[dict], instruction: str = "") -> list[dict]:
    """Expand one section; on ANY failure return the ORIGINAL blocks (never lose
    content). Returns the (possibly enriched) list of blocks."""
    from langchain_core.messages import SystemMessage, HumanMessage
    payload = json.dumps({"blocks": section}, ensure_ascii=False)
    extra = f"\n\nإضافة من المستخدم: {instruction}" if instruction else ""
    try:
        text, _label = _ai._invoke([
            SystemMessage(content=_EXPAND_SYS),
            HumanMessage(content=f"SECTION:\n{payload}{extra}"),
        ])
        spec = _ai.extract_json(text)
        blocks = spec.get("blocks") if isinstance(spec, dict) else None
        if not isinstance(blocks, list) or not blocks:
            return section
        # guard: never return fewer tables than we were given
        want_tables = sum(1 for b in section if b.get("type") == "table")
        got_tables = sum(1 for b in blocks if isinstance(b, dict) and b.get("type") == "table")
        if got_tables < want_tables:
            # model dropped a table → keep original to preserve data
            return section
        # sanitise blocks
        clean = []
        for b in blocks:
            if not isinstance(b, dict) or "type" not in b:
                continue
            clean.append(b)
        return clean or section
    except Exception:
        return section


# ── render blocks → .docx (RTL Arabic) ────────────────────────────────────────
def _is_arabic(s: str) -> bool:
    return _ai and any("؀" <= ch <= "ۿ" for ch in (s or ""))


def render_docx(blocks: list[dict], path: str, rtl: bool | None = None,
                title: str = "") -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from pathlib import Path

    if rtl is None:
        rtl = any(_is_arabic(b.get("text", "")) or
                  any(_is_arabic(str(c)) for c in b.get("header", []))
                  for b in blocks)

    doc = Document()

    def _align(par):
        if rtl:
            par.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if title:
        _align(doc.add_heading(title, level=0))

    for b in blocks:
        t = b.get("type")
        if t == "heading":
            _align(doc.add_heading(b.get("text", ""), level=min(max(b.get("level", 1), 0), 9)))
        elif t == "bullet":
            _align(doc.add_paragraph(b.get("text", ""), style="List Bullet"))
        elif t == "para":
            _align(doc.add_paragraph(b.get("text", "")))
        elif t == "table":
            header = b.get("header", []) or []
            rows = b.get("rows", []) or []
            all_rows = ([header] if header else []) + rows
            if not all_rows:
                continue
            cols = max(len(r) for r in all_rows)
            tbl = doc.add_table(rows=0, cols=cols)
            tbl.style = "Light Grid Accent 1"
            if rtl:
                tbl._tbl.tblPr.append(tbl._tbl.tblPr.makeelement(qn("w:bidiVisual"), {}))
            for i, r in enumerate(all_rows):
                cells = tbl.add_row().cells
                for j in range(cols):
                    cells[j].text = str(r[j]) if j < len(r) else ""
                    for p in cells[j].paragraphs:
                        _align(p)
                        for run in p.runs:
                            run.font.size = Pt(11)
                            if header and i == 0:
                                run.font.bold = True
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


# ── orchestrator ──────────────────────────────────────────────────────────────
def expand_document(src: str, dest: str, instruction: str = "",
                    progress=None, max_workers: int | None = None) -> dict:
    """Read `src`, expand every section with the strong model (preserving all
    tables), and write a professional RTL .docx to `dest`.

    Sections are independent, so they are expanded CONCURRENTLY (bounded by
    max_workers, default from HAYO_EXPAND_WORKERS or 6) — cutting wall-clock time
    ~6× versus sequential — while output order is preserved by section index.

    Returns stats {sections, tables_in, tables_out, blocks_in, blocks_out,
    words_in, words_out, expanded_sections}. `progress(done, n)` is called as
    each section finishes if given."""
    import os as _os
    from concurrent.futures import ThreadPoolExecutor, as_completed

    blocks = read_docx_blocks(src)
    sections = segment_sections(blocks)
    tables_in = sum(1 for b in blocks if b["type"] == "table")
    words_in = sum(len(b.get("text", "").split()) for b in blocks)

    n = len(sections)
    if max_workers is None:
        max_workers = int(_os.getenv("HAYO_EXPAND_WORKERS", "6"))
    max_workers = max(1, min(max_workers, n or 1))

    results: list[list[dict] | None] = [None] * n
    expanded = 0
    done = 0
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        futs = {pool.submit(expand_section, sec, instruction): idx
                for idx, sec in enumerate(sections)}
        for fut in as_completed(futs):
            idx = futs[fut]
            try:
                new = fut.result()
            except Exception:
                new = sections[idx]
            results[idx] = new
            if new is not sections[idx] and new != sections[idx]:
                expanded += 1
            done += 1
            if progress:
                try:
                    progress(done, n)
                except Exception:
                    pass

    out_blocks: list[dict] = []
    for idx in range(n):
        out_blocks.extend(results[idx] if results[idx] is not None else sections[idx])

    tables_out = sum(1 for b in out_blocks if isinstance(b, dict) and b.get("type") == "table")
    # absolute safety net: never emit fewer tables than the source
    if tables_out < tables_in:
        out_blocks = blocks
        tables_out = tables_in
        expanded = 0
    words_out = sum(len(b.get("text", "").split()) for b in out_blocks if isinstance(b, dict))
    path = render_docx(out_blocks, dest)
    return {
        "path": path, "sections": n,
        "tables_in": tables_in, "tables_out": tables_out,
        "blocks_in": len(blocks), "blocks_out": len(out_blocks),
        "words_in": words_in, "words_out": words_out,
        "expanded_sections": expanded,
    }
